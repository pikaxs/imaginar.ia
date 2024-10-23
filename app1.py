from flask import Flask, request, render_template, redirect, url_for, session, send_file
import os
import random
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import asyncio

app = Flask(__name__)
app.secret_key = 'patata1'

# Simulación de OpenAI API
class MockOpenAI:
    def completions_create(self, model, messages):
        # Simulamos una respuesta fija de la API de OpenAI
        return {
            "choices": [
                {
                    "message": {
                        "content": f"Este es un cuento simulado basado en el tema: {messages[-1]['content']}.\nPREGUNTAS DE COMPRENSIÓN LECTORA:\n1. ¿Cuál es el tema del cuento?\n2. ¿Quién es el personaje principal?"
                    }
                }
            ]
        }

    def chat_completions_create(self, model, messages):
        return self.completions_create(model, messages)

    def audio_speech_create(self, model, voice, input):
        # Simulamos una respuesta de audio
        return {"content": b"Simulacion de audio."}

    def images_generate(self, model, prompt, size, quality, n):
        # Simulamos una respuesta de imagen
        return {
            "data": [
                {"url": f"https://example.com/simulated_image_{random.randint(1, 100)}.png"} for _ in range(n)
            ]
        }

# Usamos la clase simulada para no gastar tokens de OpenAI
client = MockOpenAI()

# Función init_conversacion (ahora fuera de la clase)
def init_conversacion():
    if 'conversacion' not in session:
        session['conversacion'] = [
            {
                "role": "system",
                "content": """
                Eres un experto en crear cuentos infantiles para niños. Siempre debes seguir estas reglas:

                1. **Lecciones morales o éticas**: Cada cuento que crees debe enseñar una lección moral, ética o relacionada con el tema sugerido por el usuario. Estas lecciones pueden incluir valores como la honestidad, la amistad, el respeto, la seguridad en internet, entre otros. No crearás cuentos sin un componente educativo.

                2. **Temas apropiados**: Solo debes permitir temas apropiados para niños. Si el tema es inapropiado, debes rechazarlo y pedirle al usuario que reinicie el cuento con el botón de "reiniciar cuento". Los temas inapropiados incluyen violencia, lenguaje ofensivo, y cualquier contenido no adecuado para menores.

                3. **Interactividad con preguntas**: En cada cuento debes incluir preguntas interactivas en momentos clave de la historia. Debes pausar la narrativa y esperar una respuesta del usuario. Según la respuesta:
                    - Si la respuesta es correcta, continúa la historia y refuerza la enseñanza con un comentario positivo.
                    - Si la respuesta es incorrecta, corrige amablemente y proporciona la explicación adecuada.
                    - Si la respuesta no está relacionada, repite la pregunta y espera una respuesta adecuada.

                4. **Variedad en las preguntas**: Usa diferentes tipos de preguntas interactivas, como preguntas de opción múltiple, decisiones que afecten la historia, y preguntas abiertas que fomenten la creatividad del usuario.

                5. **Respetar la interactividad**: Debes siempre esperar la respuesta del usuario antes de continuar la historia.

                6. **Corrección y refuerzos positivos**: Si el usuario responde incorrectamente, corrige respetuosamente. Si responde correctamente, ofrécele un refuerzo positivo como una felicitación o motivación.

                7. **Seguridad en internet** (cuando aplique): Si el cuento trata sobre el uso de internet, refuerza la importancia de no compartir información personal, consultar a un adulto, y no hacer clic en anuncios sospechosos.

                8. **Expansión infinita de la historia**: Cada cuento debe tener la posibilidad de extenderse indefinidamente según las decisiones del usuario. Genera nuevas situaciones, personajes y lecciones según las respuestas dadas.
                """
            }
        ]

@app.route('/download_word', methods=['GET', 'POST'])
def generar_word():
    try:
        word_path = os.path.join(app.root_path, 'static', 'preguntas_comprension.docx')

        # Crear el documento Word
        doc = Document()

        # Título del documento
        title = doc.add_heading('Cuento Completo y Preguntas de Comprensión Lectora', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el título

        # Simulamos el cuento y las preguntas
        cuento = "Este es un cuento simulado."
        preguntas = "1. ¿Cuál es el tema del cuento?\n2. ¿Quién es el personaje principal?"

        # Agregar el cuento al documento
        doc.add_heading('Cuento Completo', level=2)
        cuento_paragraph = doc.add_paragraph(cuento)
        cuento_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Estilo del texto del cuento
        for run in cuento_paragraph.runs:
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x99)
            run.font.name = 'Comic Sans MS'

        # Espacio antes de las preguntas
        doc.add_paragraph('\n')

        # Agregar las preguntas de comprensión lectora
        doc.add_heading('Preguntas de Comprensión Lectora', level=2)
        for pregunta in preguntas.split("\n"):
            pregunta_paragraph = doc.add_paragraph(f'• {pregunta}')
            pregunta_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in pregunta_paragraph.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0xFF, 0x33, 0x33)
                run.font.name = 'Comic Sans MS'

        doc.save(word_path)

        return send_file(word_path, as_attachment=True, download_name='preguntas_comprension.docx')

    except Exception as e:
        print(f"Error al generar el archivo Word: {e}")
        return "Error al generar el archivo Word", 500

def generar_cuento(tema):
    init_conversacion()
    conversacion = session['conversacion']
    text = f"Simulación de cuento sobre el tema: {tema}"

    conversacion.append({"role": "user", "content": text})
    
    # Simulación de la respuesta del modelo
    completion = client.completions_create(
        model="gpt-4",
        messages=conversacion
    )
    
    aux = completion["choices"][0]["message"]["content"]
    conversacion.append({"role": "assistant", "content": aux})
    session['conversacion'] = conversacion
    return aux

async def generar_voz_async(cuento):
    # Simulamos la creación de audio
    file_path = "static/output_simulated.mp3"
    with open(file_path, "wb") as f:
        f.write(b"Simulacion de audio.")  # Audio simulado
    return file_path

async def generar_img_async(fragmento, i):
    output_folder = os.path.join(app.root_path, 'static', 'images')
    os.makedirs(output_folder, exist_ok=True)
    
    # Simulamos la generación de imágenes
    output_path = os.path.join(output_folder, f"image{i+1}_simulated.png")
    with open(output_path, 'wb') as f:
        f.write(b"Simulacion de imagen.")  # Imagen simulada
    return output_path

def generar_respuesta(user_input):
    conversacion = session.get('conversacion', [])
    conversacion.append({"role": "system", "content": "Simulamos la interacción con el usuario."})
    conversacion.append({"role": "user", "content": user_input})

    completion = client.completions_create(
        model="gpt-4",
        messages=conversacion
    )

    sop = completion["choices"][0]["message"]["content"]
    conversacion.append({"role": "assistant", "content": sop})
    
    session['conversacion'] = conversacion
    return sop    

async def generar_contenido_async(cuento):
    # Simulamos la generación de contenido de voz e imágenes
    audio_task = asyncio.create_task(generar_voz_async(cuento))
    image_tasks = [asyncio.create_task(generar_img_async(cuento, i)) for i in range(2)]

    audio_path = await audio_task
    image_paths = await asyncio.gather(*image_tasks)
    
    return audio_path, image_paths

@app.route('/', methods=['GET', 'POST'])
async def index():
    init_conversacion()
    cuento = None
    audio_url = None
    image_urls = []

    try:
        if request.method == 'POST':
            user_input = request.form.get('mic_text') or request.form.get('user_input')

            if user_input:
                cuento = generar_respuesta(user_input)
                audio_path, image_paths = await generar_contenido_async(cuento)
                audio_url = url_for('static', filename='output_simulated.mp3')
                image_urls = [url_for('static', filename=f'images/image{idx+1}_simulated.png') for idx in range(len(image_paths)) if image_paths[idx] is not None]
                return render_template('inicio.html', cuento=cuento, audio_url=audio_url, image_urls=image_urls)

            elif 'theme' in request.form:
                tema = request.form['theme']
                cuento = generar_cuento(tema)
                audio_path, image_paths = await generar_contenido_async(cuento)
                audio_url = url_for('static', filename='output_simulated.mp3')
                image_urls = [url_for('static', filename=f'images/image{idx+1}_simulated.png') for idx in range(len(image_paths)) if image_paths[idx] is not None]
                return render_template('inicio.html', cuento=cuento, audio_url=audio_url, image_urls=image_urls)

        return render_template('inicio.html', cuento=cuento, audio_url=audio_url, image_urls=image_urls)

    except Exception as e:
        print(f"Error: {e}")
        return redirect(url_for('reset'))

@app.route('/reset', methods=['POST'])
def reset():
    session.pop('conversacion', None)
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
