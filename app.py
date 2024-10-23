from flask import Flask, request, render_template, redirect, url_for, session,send_file
from openai import OpenAI  
import asyncio
from pathlib import Path
import os
import requests
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import tempfile
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

app = Flask(__name__)
app.secret_key = 'patata1'

client = OpenAI()
executor = ThreadPoolExecutor(max_workers=5)

def init_conversacion():
    if 'conversacion' not in session:
        session['conversacion'] = [
    {
        "role": "system",
        "content": """
        Eres un experto en crear cuentos infantiles para niños. Siempre debes seguir estas reglas:

        1. **Lecciones morales o éticas**: Cada cuento que crees debe enseñar una lección moral, ética o relacionada con el tema sugerido por el usuario. Estas lecciones pueden incluir valores como la honestidad, la amistad, el respeto, la seguridad en internet, entre otros. No crearás cuentos sin un componente educativo.

        2. **Temas apropiados**: Solo debes permitir temas apropiados para niños. Si el tema es inapropiado, debes rechazarlo y pedirle al usuario que reinicie el cuento con el botón de "reiniciar cuento". Los temas inapropiados incluyen violencia, lenguaje ofensivo, y cualquier contenido no adecuado para menores.

        3. **Interactividad con preguntas**: En cada cuento debes incluir preguntas interactivas en momentos clave de la historia. Debes pausar la narrativa y esperar una respuesta del usuario. Según la respuesta:
            - Si la respuesta es correcta, continúa la historia y refuerza la enseñanza con un comentario positivo.
            - Si la respuesta es incorrecta, corrige amablemente y proporciona la explicación adecuada.
            - Si la respuesta no está relacionada, repite la pregunta y espera una respuesta adecuada.

        4. **Variedad en las preguntas**: Usa diferentes tipos de preguntas interactivas, como preguntas de opción múltiple, decisiones que afecten la historia, y preguntas abiertas que fomenten la creatividad del usuario.

        5. **Respetar la interactividad**: Debes siempre esperar la respuesta del usuario antes de continuar la historia.

        6. **Corrección y refuerzos positivos**: Si el usuario responde incorrectamente, corrige respetuosamente. Si responde correctamente, ofrécele un refuerzo positivo como una felicitación o motivación.

        7. **Seguridad en internet** (cuando aplique): Si el cuento trata sobre el uso de internet, refuerza la importancia de no compartir información personal, consultar a un adulto, y no hacer clic en anuncios sospechosos.

        8. **Expansión infinita de la historia**: Cada cuento debe tener la posibilidad de extenderse indefinidamente según las decisiones del usuario. Genera nuevas situaciones, personajes y lecciones según las respuestas dadas.
        """
    }
]

@app.route('/download_word', methods=['GET', 'POST'])
def generar_word():
    try:
        # Ruta del archivo en la carpeta 'static'
        word_path = os.path.join(app.root_path, 'static', 'preguntas_comprension.docx')

        # Crear el documento Word
        doc = Document()

        # Título del documento
        title = doc.add_heading('Cuento Completo y Preguntas de Comprensión Lectora', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centrar el título

        # Obtener la conversación desde la sesión
        conversacion = session.get('conversacion', [])
        text = ("quiero que escribas el cuento completo si es que te conteste la pregunta interactiva o no "
                "escribe todo el cuento completo sin que me digas claro y confirmando la orden. No únicamente "
                "quiero que escribas el cuento completo sin importar si te conteste o no. Además, quiero que "
                "generes, con respecto al cuento y/o respuestas que tuvimos, preguntas de comprensión lectora. "
                "Por lo menos 12 preguntas de comprensión lectora sin respuesta. Dale una separación, es decir, "
                "el cuento completo después un espacio en blanco y después las preguntas.")
        conversacion.append({"role": "user", "content": text})

        # Llamada a OpenAI API para generar el cuento y las preguntas
        completion = client.chat.completions.create(
            model="gpt-4",
            messages=conversacion
        )

        # Verificamos si OpenAI retornó una respuesta válida
        if not completion or not completion.choices or len(completion.choices) == 0:
            raise ValueError("No se pudo obtener una respuesta válida de OpenAI")
        
        sop = completion.choices[0].message.content

        # Dividir el contenido en cuento y preguntas
        if "PREGUNTAS DE COMPRENSIÓN LECTORA:" in sop:
            cuento, preguntas = sop.split("PREGUNTAS DE COMPRENSIÓN LECTORA:", 1)
        else:
            cuento = sop
            preguntas = ""

        # Agregar el cuento al documento
        doc.add_heading('Cuento Completo', level=2)
        cuento_paragraph = doc.add_paragraph(cuento.strip())
        cuento_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Estilo del texto del cuento (más amigable para niños)
        for run in cuento_paragraph.runs:
            run.font.size = Pt(12)  # Tamaño de letra
            run.font.color.rgb = RGBColor(0x00, 0x33, 0x99)  # Azul oscuro
            run.font.name = 'Comic Sans MS'  # Fuente amigable para niños

        # Espacio antes de las preguntas
        doc.add_paragraph('\n')

        # Agregar las preguntas de comprensión lectora
        if preguntas.strip():
            doc.add_heading('Preguntas de Comprensión Lectora', level=2)

            for pregunta in preguntas.strip().split("\n"):
                if pregunta.strip():
                    pregunta_paragraph = doc.add_paragraph(f'• {pregunta.strip()}')
                    pregunta_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in pregunta_paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0xFF, 0x33, 0x33)  # Rojo amigable
                        run.font.name = 'Comic Sans MS'

        # Guardar el documento en la ruta especificada
        doc.save(word_path)

        # Devolver el archivo generado
        return send_file(word_path, as_attachment=True, download_name='preguntas_comprension.docx')

    except Exception as e:
        print(f"Error al generar el archivo Word: {e}")
        return "Error al generar el archivo Word", 500
    
def generar_cuento(tema):
    init_conversacion()
    conversacion = session['conversacion']
    text = f"a continuacion te dare una instruccion para crear un cuento sigue las reglas y crea el cuento respecto a lo que te diga y las reglas establecidad, {tema}"

    conversacion.append({"role": "user", "content": text})

    completion = client.chat.completions.create(
        model="gpt-4",
        messages=conversacion
    )
    
    aux = completion.choices[0].message.content
    conversacion.append({"role": "assistant", "content": aux})
    session['conversacion'] = conversacion  # Actualizar la conversación en la sesión
    return aux

async def generar_voz_async(cuento):
    response = client.audio.speech.create(
        model="tts-1",
        voice="alloy",
        input=cuento
    )
    
    # Guardar el archivo de voz en la carpeta static
    file_path = "static/output.mp3"
    
    with open(file_path, "wb") as f:
        f.write(response.content)
    
    return file_path

async def generar_img_async(fragmento, i):
    output_folder = os.path.join(app.root_path, 'static', 'images')
    os.makedirs(output_folder, exist_ok=True)

    response = client.images.generate(
        model="dall-e-3",
        prompt=fragmento,
        size="1024x1024",
        quality="standard",
        n=1,
    )

    image_url = response.data[0].url
    image_name = f"image{i+1}.png"
    output_path = os.path.join(output_folder, image_name)

    image_response = requests.get(image_url)
    
    if image_response.status_code == 200:
        with open(output_path, 'wb') as f:
            f.write(image_response.content)
        print(f"Imagen {i+1} guardada en: {output_path}")
        return output_path
    else:
        print(f"Error al descargar la imagen {i+1}.")
        return None
    
def generar_respuesta(user_input):
    conversacion = session.get('conversacion', [])
    conversacion.append({"role": "system", "content": "Espera la respuesta del cuento. Si es incorrecta, responde adecuadamente mientras que si es una respuesta no relacionada al cuento escribe nuevamente la pregunta interactiva y esperas la respuesta"})
    conversacion.append({"role": "user", "content": user_input})

    completion = client.chat.completions.create(
        model="gpt-4",
        messages=conversacion
    )

    sop = completion.choices[0].message.content
    conversacion.append({"role": "assistant", "content": sop})
    
    session['conversacion'] = conversacion  # Actualizar la conversación en la sesión
    return sop    

async def generar_contenido_async(cuento):
    # Generamos voz
    audio_task = asyncio.create_task(generar_voz_async(cuento))
    
    # Generamos imágenes
    image_tasks = [asyncio.create_task(generar_img_async(cuento, i)) for i in range(2)]
    
    # Esperar a que se completen todas las tareas
    audio_path = await audio_task
    image_paths = await asyncio.gather(*image_tasks)
    
    return audio_path, image_paths

@app.route('/', methods=['GET', 'POST'])
async def index():
    init_conversacion()
    cuento = None
    audio_url = None
    image_urls = []

    try:
        if request.method == 'POST':
            user_input = request.form.get('mic_text') or request.form.get('user_input')

            if user_input:
                cuento = generar_respuesta(user_input)
                audio_path, image_paths = await generar_contenido_async(cuento)
                audio_url = url_for('static', filename='output.mp3')
                image_urls = [url_for('static', filename=f'images/image{idx+1}.png') for idx in range(len(image_paths)) if image_paths[idx] is not None]
                return render_template('inicio.html', cuento=cuento, audio_url=audio_url, image_urls=image_urls)

            # Verificamos solo el campo 'theme'
            elif 'theme' in request.form:
                tema = request.form['theme']  # Solo capturamos el campo 'theme'
                cuento = generar_cuento(tema)  # Puedes especificar un estilo fijo si lo deseas
                audio_path, image_paths = await generar_contenido_async(cuento)
                audio_url = url_for('static', filename='output.mp3')
                image_urls = [url_for('static', filename=f'images/image{idx+1}.png') for idx in range(len(image_paths)) if image_paths[idx] is not None]
                return render_template('inicio.html', cuento=cuento, audio_url=audio_url, image_urls=image_urls)

        return render_template('inicio.html', cuento=cuento, audio_url=audio_url, image_urls=image_urls)


    except Exception as e:
        print(f"Error: {e}")
        return redirect(url_for('reset'))

@app.route('/reset', methods=['POST'])
def reset():
    session.pop('conversacion', None)  # Eliminar la conversación de la sesión
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
